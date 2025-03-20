import os
import json
import re
import sys
import time
import nltk
import spacy
import pdfplumber
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from groq import Groq
from dotenv import load_dotenv
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from spacy.matcher import PhraseMatcher
from tenacity import retry, stop_after_attempt, wait_fixed

# Check dependencies
required_libs = ['pdfplumber', 'docx', 'sklearn', 'groq', 'dotenv', 'nltk', 'spacy']
for lib in required_libs:
    try:
        __import__(lib)
    except ImportError:
        print(f"Error: {lib} not installed. Install with 'pip install {lib}'")
        sys.exit(1)

# Load spaCy with parser enabled
nlp = spacy.load("en_core_web_sm")

# Download NLTK resources
try:
    nltk.download("punkt", quiet=True)
    nltk.download("stopwords", quiet=True)
except Exception as e:
    print(f"Error downloading NLTK resources: {e}")
    sys.exit(1)
STOPWORDS = set(stopwords.words("english"))

# Load API key
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    GROQ_API_KEY = input("Enter Groq API key: ").strip()
    if not GROQ_API_KEY:
        raise ValueError("Please set GROQ_API_KEY in .env or provide it manually")
client = Groq(api_key=GROQ_API_KEY)

def generate_resume(user_info):
    if not user_info or not isinstance(user_info, str):
        print("Error: user_info must be a non-empty string")
        return None
    prompt = f"Create a resume for the following information: {user_info}"
    try:
        response = client.chat.completions.create(
            model="mixtral-8x7b-32768",
            messages=[{"role": "system", "content": "You are a helpful assistant."},
                      {"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"Error generating resume: {e}")
        return None

class Resume:
    def __init__(self, file_path):
        self.file_path = file_path

    def load_resume(self):
        if not os.path.exists(self.file_path):
            print(f"File not found: {self.file_path}")
            return ""
        if not self.file_path.endswith(('.pdf', '.docx', '.txt')):
            print(f"Unsupported file format: {self.file_path}")
            return ""
        try:
            if self.file_path.endswith('.pdf'):
                with pdfplumber.open(self.file_path) as pdf:
                    return "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())
            elif self.file_path.endswith('.docx'):
                doc = Document(self.file_path)
                return "\n".join(para.text for para in doc.paragraphs)
            elif self.file_path.endswith('.txt'):
                with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    return file.read().strip()
        except Exception as e:
            print(f"Error loading resume: {e}")
            return ""

JOB_TITLE_KEYWORDS = {"data scientist", "machine learning engineer", "software engineer", "backend developer", "frontend developer", "full stack developer", "data analyst", "business analyst", "project manager", "product manager", "AI engineer", "deep learning engineer", "computer vision engineer", "NLP engineer", "big data engineer", "cloud engineer", "devops engineer", "cybersecurity analyst", "network engineer"}
DEGREE_KEYWORDS = {"bachelor", "master", "phd", "associate", "b.tech", "m.tech", "msc"}

class Skills:
    @staticmethod
    def preprocess_text(text):
        start_time = time.time()
        tokens = [token.lower() for token in word_tokenize(text) if token.lower() not in STOPWORDS and len(token) > 2]
        result = " ".join(tokens)
        print(f"Preprocess time: {time.time() - start_time:.2f}s")
        return result

    @staticmethod
    def clean_skill(skill):
        skill = re.sub(r'^\d+\.\s*', '', skill.strip().lower())
        skill = re.sub(r'^(here is).*', '', skill)
        if re.match(r'^(delivering|documenting|collaborating|staying|ensuring|requirements|tasks|analysis$)', skill):
            return None
        return skill if len(skill.split()) <= 4 else None

    @staticmethod
    def get_seed_skills(job_title):
        base_skills = [
            "python", "java", "sql", "javascript", "communication", "leadership", "excel", "aws", "docker", "git"
        ]
        role_specific = {
            "frontend": ["html", "css", "jquery", "react", "angular", "vue", "bootstrap", "sass", "webpack", "photoshop", "adobe suite", "seo"],
            "backend": ["rest apis", "nodejs", "ruby", "php", "mongodb", "redis", "postgresql", "microservices"],
            "data": ["machine learning", "data analysis", "pandas", "numpy", "matplotlib", "seaborn", "tableau", "power bi", "nlp", "deep learning", "tensorflow", "pytorch", "opencv"],
            "devops": ["kubernetes", "jenkins", "ansible", "terraform", "ci/cd", "linux"],
            "full stack": ["html", "css", "javascript", "react", "rest apis", "mongodb"]
        }
        job_title_lower = job_title.lower()
        for role, skills in role_specific.items():
            if role in job_title_lower:
                return base_skills + skills
        return base_skills

    @staticmethod
    @retry(stop=stop_after_attempt(3), wait=wait_fixed(2))
    def extract_skills(text, job_title="", is_job_description=False, use_groq=True):
        start_time = time.time()
        processed_text = Skills.preprocess_text(text)
        doc = nlp(processed_text)
        phrase_matcher = PhraseMatcher(nlp.vocab, attr="LOWER")
        seed_skills = Skills.get_seed_skills(job_title)
        phrase_matcher.add("SEED_SKILLS", [nlp(skill) for skill in seed_skills])
        skills = set()
        matches = phrase_matcher(doc)
        for match_id, start, end in matches:
            skills.add(doc[start:end].text.lower())
        for chunk in doc.noun_chunks:
            cleaned = Skills.clean_skill(chunk.text)
            if cleaned:
                skills.add(cleaned)
        if is_job_description and use_groq and len(skills) < 10:
            groq_start = time.time()
            prompt = f"Extract a concise list of technical and soft skills for a {job_title} from this job description: {processed_text[:2000]}"
            response = client.chat.completions.create(
                model="mixtral-8x7b-32768",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=150
            )
            groq_skills = {Skills.clean_skill(skill) for skill in response.choices[0].message.content.split("\n") if skill.strip()}
            skills.update(skill for skill in groq_skills if skill)
            print(f"Groq time: {time.time() - groq_start:.2f}s")
        skills = {skill for skill in skills if skill and not re.match(r'^(experience|knowledge|ability|skills)$', skill)}
        print(f"Extract skills time: {time.time() - start_time:.2f}s")
        return skills

    @staticmethod
    def extract_resume_skills(resume_content, job_title=""):
        return Skills.extract_skills(resume_content, job_title, is_job_description=False, use_groq=False)

    @staticmethod
    def categorize_skills_with_groq(missing_skills):
        if not missing_skills:
            return {"hard_skills": [], "soft_skills": [], "other_skills": []}
        prompt = f"Categorize these skills into hard, soft, and other: {', '.join(missing_skills[:20])}"
        try:
            response = client.chat.completions.create(
                model="mixtral-8x7b-32768",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=50
            )
            content = response.choices[0].message.content.strip()
            return json.loads(content)
        except json.JSONDecodeError:
            print(f"Error parsing Groq response: {content}")
            return {"hard_skills": [], "soft_skills": [], "other_skills": []}
        except Exception as e:
            print(f"Error categorizing skills: {e}")
            return {"hard_skills": [], "soft_skills": [], "other_skills": []}

    @staticmethod
    def get_relevant_skills(job_title, job_description):
        return Skills.extract_skills(job_description, job_title, is_job_description=True, use_groq=True)

    @staticmethod
    def find_missing_skills(resume_content, job_title, job_description):
        job_skills = Skills.get_relevant_skills(job_title, job_description)
        resume_skills = Skills.extract_resume_skills(resume_content, job_title)
        print(f"Job Skills: {sorted(job_skills)}")
        print(f"Resume Skills: {sorted(resume_skills)}")
        return sorted(list(job_skills - resume_skills))

    @staticmethod
    def extract_job_title(job_description):
        processed_text = Skills.preprocess_text(job_description)
        for keyword in JOB_TITLE_KEYWORDS:
            if re.search(rf"\b{keyword}\b", processed_text, re.IGNORECASE):
                return keyword
        return "Unknown"

    @staticmethod
    def check_degree_match(resume_content):
        processed_text = Skills.preprocess_text(resume_content)
        return any(degree in processed_text for degree in DEGREE_KEYWORDS)

    @staticmethod
    def calculate_ats_score(resume_content, job_description):
        start_time = time.time()
        resume_processed = Skills.preprocess_text(resume_content)
        job_processed = Skills.preprocess_text(job_description)
        vectorizer = TfidfVectorizer().fit_transform([resume_processed, job_processed])
        similarity = cosine_similarity(vectorizer[0], vectorizer[1])
        print(f"ATS score time: {time.time() - start_time:.2f}s")
        return round(similarity[0][0] * 100, 2)

def rewrite_resume(resume_content, job_description):
    prompt = f"""
    Rewrite the following resume to better match the job description while keeping it ATS-friendly.

    Resume Content:
    {resume_content[:4000]}

    Job Description:
    {job_description[:4000]}

    Ensure the rewritten resume:
    - Uses relevant keywords from the job description
    - Highlights missing skills naturally
    - Uses professional and concise language
    - Maintains proper resume formatting with clear sections:
      - 'Work Experience' for jobs and internships only
      - 'Education' for academic degrees and institutions
      - Other sections (e.g., Summary, Projects, Skills) as appropriate
    - Does not mix education details into 'Work Experience' or vice versa
    - Excludes any analysis, summary of improvements, or notes about limitations (e.g., character limits)

    Provide only the improved resume content in a structured format with section headers marked by **Header Name**. Do not include any additional explanations or comments beyond the resume itself.
    """
    try:
        response = client.chat.completions.create(
            model="mixtral-8x7b-32768",
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"Error generating rewritten resume: {e}")
        return "Error generating rewritten resume."

class Analyze:
    @staticmethod
    def analyze_resume(resume_content, job_title, job_description):
        start_time = time.time()
        word_count = len(word_tokenize(resume_content))
        word_count_message = (
            "Your resume is within the recommended range of 500-700 words."
            if 500 <= word_count <= 700 else "Consider adjusting the content for better readability."
        )
        missing_skills = Skills.find_missing_skills(resume_content, job_title, job_description)
        insights = {
            "ats_score": Skills.calculate_ats_score(resume_content, job_description),
            "missing_skills": missing_skills,
            "categorized_skills": Skills.categorize_skills_with_groq(missing_skills),
            "degree_match": Skills.check_degree_match(resume_content),
            "word_count": word_count,
            "word_count_message": word_count_message,
            "rewritten_resume": rewrite_resume(resume_content, job_description)
        }
        print(f"Analyze time: {time.time() - start_time:.2f}s")
        return insights

chat_history = [{"role": "system", "content": "You are an AI assistant helping with resume analysis and ATS optimization. Provide concise, actionable advice."}]

def chat_with_assistant(message, resume_content="", job_description="", insights=None):
    chat_history.append({"role": "user", "content": message})
    if resume_content and not any("resume:" in msg["content"] for msg in chat_history):
        chat_history.append({"role": "user", "content": f"Resume:\n{resume_content}"})
    if job_description and not any("job description:" in msg["content"] for msg in chat_history):
        chat_history.append({"role": "user", "content": f"Job Description:\n{job_description}"})
    if insights and not any("insights:" in msg["content"] for msg in chat_history):
        chat_history.append({"role": "user", "content": f"Insights:\n{json.dumps(insights, indent=2)}"})
    if len(chat_history) > 10:
        chat_history[:] = chat_history[-10:]
    try:
        response = client.chat.completions.create(
            model="mixtral-8x7b-32768",
            messages=chat_history
        )
        assistant_reply = response.choices[0].message.content
        chat_history.append({"role": "assistant", "content": assistant_reply})
        return assistant_reply
    except Exception as e:
        return f"Error in chat: {e}"

def save_rewritten_resume(content, output_path="rewritten_resume.txt"):
    try:
        if os.path.exists(output_path):
            print(f"Warning: Overwriting {output_path}")
        formatted_content = ""
        lines = content.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                formatted_content += "\n"
            elif line.startswith("**") and line.endswith("**"):
                formatted_content += f"\n{line[2:-2]}\n"
            elif line.startswith("- "):
                formatted_content += f"  - {line[2:]}\n"
            else:
                while len(line) > 80:
                    split_point = line.rfind(" ", 0, 80)
                    if split_point == -1:
                        split_point = 80
                    formatted_content += f"{line[:split_point]}\n"
                    line = line[split_point:].strip()
                formatted_content += f"{line}\n"
        with open(output_path, "w", encoding='utf-8') as file:
            file.write(formatted_content.strip())
        print(f"✅ Rewritten resume saved to: {output_path}")
    except Exception as e:
        print(f"Error saving rewritten resume: {e}")

def fill_resume_in_word(content, output_path="rewritten_resume.docx"):
    try:
        doc = Document()
        
        # Set default font
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal'].font.size = Pt(10)

        # Parse and format content
        lines = content.split("\n")
        header_keywords = {"summary", "work experience", "education", "projects", "technical skills", "soft skills", "certifications"}
        link_pattern = re.compile(r'\[(.*?)\]\(.*?\)')

        for line in lines:
            line = line.strip()
            line = link_pattern.sub(r'\1', line)  # Replace links with text only
            
            if not line:
                doc.add_paragraph("")  # Add empty line
            elif (line.startswith("**") and line.endswith("**")) or line.lower().strip() in header_keywords:
                header_text = line[2:-2] if line.startswith("**") else line
                p = doc.add_paragraph()
                run = p.add_run(header_text)
                run.bold = True
                run.font.size = Pt(12)
            elif line.startswith("- "):
                p = doc.add_paragraph(style='ListBullet')
                p.add_run(line[2:]).font.size = Pt(10)
            else:
                doc.add_paragraph(line)

        # Save the document temporarily
        doc.save(output_path)
        print(f"✅ Temporary Word document created at: {output_path}")

        # Open the document using the default application
        if sys.platform == "win32":
            os.startfile(output_path)  # Windows
        elif sys.platform == "darwin":
            os.system(f"open {output_path}")  # macOS
        else:
            os.system(f"xdg-open {output_path}")  # Linux
        print("✅ Word document opened for review. Please save it manually after editing.")
        
    except Exception as e:
        print(f"Error creating or opening Word document: {e}")

def main():
    resume_path = input("Enter your resume file path (PDF, DOCX, or TXT): ").strip()
    job_title = input("Enter the job title: ").strip()
    
    print("\nEnter job description (type 'END' on a new line to finish):")
    job_description_lines = []
    while True:
        line = input()
        if line.strip().upper() == "END":
            break
        job_description_lines.append(line)
    job_description = "\n".join(job_description_lines).strip()
    
    if not all([resume_path, job_title, job_description]):
        print("Error: All inputs are required.")
        return
    
    resume_obj = Resume(resume_path)
    resume_content = resume_obj.load_resume()
    if not resume_content:
        print("Error: Could not extract content from the resume.")
        return
    
    insights = Analyze.analyze_resume(resume_content, job_title, job_description)
    print("\n🔹 Resume Analysis Insights 🔹")
    print(json.dumps(insights, indent=4))

    print("\nResume Chat Assistant | Type 'exit' to quit.\n")
    while True:
        user_input = input("You: ")
        if user_input.lower() == "exit":
            print("Goodbye!")
            break
        response = chat_with_assistant(user_input, resume_content, job_description, insights)
        print(f"Assistant: {response}\n")
    
    save_rewritten_resume(insights["rewritten_resume"])
    fill_resume_in_word(insights["rewritten_resume"])

if __name__ == "__main__":
    main()